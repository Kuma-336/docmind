# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_font(run, size=11):
    run.font.name = "微软雅黑"
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    colors = {0: RGBColor(0x1F, 0x49, 0x7D), 1: RGBColor(0x1F, 0x49, 0x7D), 2: RGBColor(0x2E, 0x75, 0xB6)}
    for run in p.runs:
        set_font(run, 14 if level <= 1 else 12)
        run.font.color.rgb = colors.get(level, RGBColor(0x2E, 0x75, 0xB6))
    return p


def add_para(doc, text, size=11, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run, size)
    return p


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)


def init_table(doc, headers, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, (cell, txt) in enumerate(zip(table.rows[0].cells, headers)):
        cell.text = txt
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        set_font(run, 10.5)
        if col_widths:
            cell.width = Inches(col_widths[i])
    return table


def add_row(table, values, bold_first=False):
    row = table.add_row()
    for i, (cell, txt) in enumerate(zip(row.cells, values)):
        cell.text = txt
        run = cell.paragraphs[0].runs[0]
        set_font(run, 10.5)
        if bold_first and i == 0:
            run.bold = True


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run, 11)


def build():
    doc = Document()

    # 默认样式
    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal.font.size = Pt(11)

    # ── 封面标题 ──────────────────────────────────────────
    title_p = doc.add_heading("DocMind 阶段二完成报告", 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_p.runs:
        set_font(run, 18)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph("RAG 核心模块实现报告 · DocMind v0.1.0")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        set_font(run, 11)

    doc.add_paragraph()

    # ── 一、项目概述 ──────────────────────────────────────
    add_heading(doc, "一、项目概述", 1)
    add_para(
        doc,
        "DocMind 是基于 FastAPI + LangChain + LangGraph + ChromaDB 的多 Agent 文档问答系统。"
        "阶段二完成了 RAG 核心模块的全部实现，替换了阶段一的占位代码，并通过了端到端验证。"
        "技术选型：DeepSeek Chat API 作为 LLM，本地 BAAI/bge-small-zh-v1.5 作为 Embedding 模型。",
    )

    doc.add_paragraph()

    # ── 二、修改/新增文件清单 ─────────────────────────────
    add_heading(doc, "二、修改 / 新增文件清单", 1)

    t1 = init_table(doc, ["文件路径", "操作 & 说明"], [2.5, 4.0])
    for row in [
        ("app/rag/loader.py",            "重写 — DocumentLoader 类，支持 pdf/txt/md，自动补全 metadata"),
        ("app/rag/splitter.py",          "重写 — TextSplitter 类，输出 chunk_index / chunk_total"),
        ("app/rag/embedder.py",          "重写 — EmbeddingManager 单例，本地 BAAI/bge-small-zh-v1.5"),
        ("app/rag/retriever.py",         "重写 — VectorStoreManager 单例，langchain-chroma，delete_by_file_id"),
        ("app/api/routes/documents.py",  "重写 — upload/list/delete 完整业务逻辑，含错误清理"),
        ("app/agents/rag_agent.py",      "重写 — rag_node async 函数（LangGraph 节点）"),
        ("app/config.py",                "更新 — 新增 OPENAI_BASE_URL、EMBEDDING_MODEL 字段"),
        ("requirements.txt",             "更新 — 新增 langchain-chroma、sentence-transformers、langchain-huggingface"),
        ("scripts/test_rag_pipeline.py", "新建 — 端到端 pipeline 验证脚本"),
        (".env",                         "更新 — DeepSeek base_url、本地 embedding 模型名"),
    ]:
        add_row(t1, row, bold_first=True)

    doc.add_paragraph()

    # ── 三、核心模块说明 ──────────────────────────────────
    add_heading(doc, "三、核心模块说明", 1)

    add_heading(doc, "3.1  DocumentLoader（loader.py）", 2)
    add_para(
        doc,
        "支持 .pdf（PyPDFLoader）、.txt / .md（TextLoader）三种格式；不支持的格式抛出 ValueError "
        "并列出支持列表。每个 Document 的 metadata 自动写入以下字段：",
    )
    for item in [
        "file_name — 原始文件名",
        "file_path — 文件绝对路径",
        "file_type — 文件扩展名（不含点号）",
        "loaded_at — UTC ISO 8601 时间戳",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.2  TextSplitter（splitter.py）", 2)
    add_para(
        doc,
        "使用 RecursiveCharacterTextSplitter，chunk_size / chunk_overlap 从 get_settings() 读取"
        "（默认 1000 / 200）。每个 chunk 继承父 Document 的全部 metadata，并额外追加 "
        "chunk_index（从 0 起）和 chunk_total（该文档总 chunk 数）。",
    )

    add_heading(doc, "3.3  EmbeddingManager（embedder.py）", 2)
    add_para(
        doc,
        "使用本地 HuggingFace 模型 BAAI/bge-small-zh-v1.5（约 90 MB，中英双语，支持归一化向量）。"
        "首次运行自动下载并缓存到 ~/.cache/huggingface，后续离线可用。"
        "采用单例模式（_instance + get_embedder() 工厂函数），避免重复加载模型占用内存。",
    )

    add_heading(doc, "3.4  VectorStoreManager（retriever.py）", 2)
    add_para(
        doc,
        "基于 langchain-chroma 管理 ChromaDB 持久化存储，collection 固定为 \"docmind\"。提供三个方法：",
    )
    for item in [
        "add_documents(documents, file_id) — 批量写入 chunk，为每个 chunk 注入 file_id，返回写入数量",
        "retrieve(query, top_k) — 调用 similarity_search 检索，top_k 默认读取 RETRIEVER_TOP_K 配置",
        "delete_by_file_id(file_id) — 通过 chromadb where 过滤查询 ID，再批量删除对应 chunk",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.5  文档管理 API（documents.py）", 2)
    t_api = init_table(doc, ["方法", "路径", "功能说明"], [0.8, 1.5, 4.2])
    for row in [
        ("POST",   "/upload",     "保存文件 → load → split → 向量化存储，失败时自动清理已保存文件，返回 chunk_count"),
        ("GET",    "/list",       "扫描 UPLOAD_DIR，返回 file_id / filename / size / uploaded_at"),
        ("DELETE", "/{file_id}",  "删除磁盘文件 + Chroma 向量，两步均成功才返回 200"),
    ]:
        add_row(t_api, row)

    add_heading(doc, "3.6  rag_node（rag_agent.py）", 2)
    add_para(
        doc,
        "标准 LangGraph 异步节点（async def rag_node(state: AgentState) -> dict）。"
        "不直接修改传入 state，返回新 dict。检索结果以 \"---\" 分隔拼接为字符串写入 rag_results；"
        "来源文件名去重后写入 sources；\"rag_agent\" 追加到 agent_path。"
        "无检索结果时 rag_results 写入友好提示文字。",
    )

    doc.add_paragraph()

    # ── 四、验证结果 ──────────────────────────────────────
    add_heading(doc, "四、验证结果", 1)

    add_heading(doc, "4.1  test_rag_pipeline.py 输出", 2)
    add_code_block(doc, [
        "[1] 已创建测试文件: data/uploads/test_doc.txt",
        "[2] 加载完成，共 1 个文档段",
        "[3] 切分完成，共 1 个 chunk",
        "    chunk[0] 长度=749，index=0/1",
        "[4] 写入向量库完成，共 1 个 chunk，file_id=test-pipeline-001",
        "[5] 检索问题「什么是机器学习？」→ 1 条结果",
        '    metadata: {file_name: "test_doc.txt", chunk_index: 0, chunk_total: 1, ...}',
        "[5] 检索问题「深度学习有什么应用？」→ 1 条结果",
        "RAG pipeline 验证完成，共检索到 2 条结果",
    ])

    add_heading(doc, "4.2  上传接口返回（POST /api/v1/documents/upload）", 2)
    add_code_block(doc, [
        '{',
        '  "file_id": "178f003b-99e4-4afb-baf7-6f94e9a49dab",',
        '  "filename": "test_doc.txt",',
        '  "chunk_count": 1,',
        '  "status": "success"',
        '}',
    ])

    add_heading(doc, "4.3  文件列表返回（GET /api/v1/documents/list）", 2)
    add_code_block(doc, [
        '[',
        '  {',
        '    "file_id": "178f003b-99e4-4afb-baf7-6f94e9a49dab",',
        '    "filename": "test_doc.txt",',
        '    "size": 2155,',
        '    "uploaded_at": "2026-06-14T18:24:53.587496+00:00"',
        '  }',
        ']',
    ])

    add_heading(doc, "4.4  删除接口返回（DELETE /api/v1/documents/{file_id}）", 2)
    add_code_block(doc, [
        '{',
        '  "file_id": "178f003b-99e4-4afb-baf7-6f94e9a49dab",',
        '  "status": "deleted"',
        '}',
    ])

    doc.add_paragraph()

    # ── 五、问题与解决 ────────────────────────────────────
    add_heading(doc, "五、遇到的问题与解决方式", 1)

    t_issue = init_table(doc, ["问题", "解决方式"], [3.0, 3.5])
    for row in [
        ("DeepSeek 无 Embedding API，OpenAI 占位 Key 返回 401",
         "改用本地 BAAI/bge-small-zh-v1.5（langchain-huggingface），DeepSeek 仅用于 Chat LLM"),
        ("langchain_community.vectorstores.Chroma 已弃用",
         "安装 langchain-chroma 1.1.0 并更新 import"),
        ("HuggingFaceEmbeddings 弃用警告",
         "安装 langchain-huggingface 1.2.2 并更新 import"),
        ("Windows 下 HuggingFace 缓存不支持符号链接",
         "非阻塞 Warning，功能正常，可开启开发者模式或以管理员运行解决"),
        ("PowerShell 中文日志乱码",
         "终端编码问题，数据本身正确；可设 PYTHONIOENCODING=utf-8 环境变量解决"),
    ]:
        add_row(t_issue, row)

    doc.add_paragraph()

    # ── 六、阶段三就绪 ────────────────────────────────────
    add_heading(doc, "六、阶段三（LangGraph 多 Agent 编排）就绪标志", 1)

    for item in [
        "rag_node 已实现为标准 LangGraph 异步节点，接口兼容 StateGraph.add_node()",
        "AgentState 中 rag_results、sources、agent_path 字段由 rag_node 正确填充，格式符合后续节点预期",
        "DeepSeek Chat API 已通过 OPENAI_BASE_URL + OPENAI_API_KEY 配置，ChatOpenAI 可直接使用",
        "ChromaDB 持久化目录 ./data/chroma 已就绪，数据跨进程保留",
        "待实现：app/graph/builder.py（图编排）、search_agent.py、summarizer_agent.py、supervisor.py",
    ]:
        add_bullet(doc, item)

    out = r"C:\Users\kuma\Desktop\找工\DocMind阶段二完成报告.docx"
    doc.save(out)
    print("已保存:", out)


if __name__ == "__main__":
    build()
